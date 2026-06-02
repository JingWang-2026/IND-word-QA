from __future__ import annotations

import json
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from docx.document import Document as PythonDocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from lxml import etree
from sqlmodel import Session, delete

from app.models.parsed_content import ParsedBlock, ParsedTable
from app.schemas.parsed_content import (
    ParsedComment,
    ParsedDocument,
    ParsedHeading,
    ParsedParagraph,
    ParsedTableData,
    ParsedTextBlock,
    TrackedChange,
)

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
HEADING_NUMBER_RE = re.compile(r"^\s*(?P<number>\d+(?:\.\d+)*)(?:\.|\s)\s*(?P<title>.+)$")


@dataclass
class SectionContext:
    number: str | None = None
    title: str | None = None


class DocxParserService:
    def parse(self, document_id: str, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        doc = DocxDocument(path)
        parsed = ParsedDocument(document_id=document_id, metadata=self._metadata(doc))
        context = SectionContext()

        previous_paragraphs: list[str] = []
        table_index = 0
        paragraph_index = 0

        for block in self._iter_body_blocks(doc):
            if isinstance(block, Paragraph):
                text = self._clean_text(block.text)
                style_name = block.style.name if block.style is not None else None
                if not text:
                    paragraph_index += 1
                    continue

                heading = self._maybe_heading(text, style_name, paragraph_index)
                if heading is not None:
                    context.number = heading.section_number
                    context.title = heading.section_title
                    parsed.headings.append(heading)

                parsed.paragraphs.append(
                    ParsedParagraph(
                        text=text,
                        paragraph_index=paragraph_index,
                        index=paragraph_index,
                        style_name=style_name,
                        section_number=context.number,
                        section_title=context.title,
                    )
                )
                previous_paragraphs.append(text)
                paragraph_index += 1
            elif isinstance(block, DocxTable):
                rows = [[self._clean_text(cell.text) for cell in row.cells] for row in block.rows]
                caption = self._find_caption(previous_paragraphs[-3:])
                parsed.tables.append(
                    ParsedTableData(
                        table_index=table_index,
                        rows=rows,
                        markdown=self._to_markdown(rows),
                        caption=caption,
                        section_number=context.number,
                        section_title=context.title,
                    )
                )
                table_index += 1

        parsed.headers, parsed.footers = self._headers_and_footers(doc)
        parsed.comments = self._comments(path)
        parsed.tracked_changes = self._tracked_changes(path)
        parsed.hidden_text = self._hidden_text(path)
        return parsed

    def save_to_database(self, session: Session, parsed: ParsedDocument) -> None:
        session.exec(delete(ParsedBlock).where(ParsedBlock.document_id == parsed.document_id))
        session.exec(delete(ParsedTable).where(ParsedTable.document_id == parsed.document_id))

        for paragraph in parsed.paragraphs:
            self._add_block(
                session,
                parsed.document_id,
                block_type="paragraph",
                text=paragraph.text,
                paragraph_index=paragraph.paragraph_index,
                section_number=paragraph.section_number,
                section_title=paragraph.section_title,
                style_name=paragraph.style_name,
            )

        for heading in parsed.headings:
            self._add_block(
                session,
                parsed.document_id,
                block_type="heading",
                text=heading.text,
                paragraph_index=heading.paragraph_index,
                section_number=heading.section_number,
                section_title=heading.section_title,
                style_name=heading.style_name,
                metadata={"level": heading.level},
            )

        for block in parsed.headers:
            self._add_block(
                session,
                parsed.document_id,
                block_type="header",
                text=block.text,
                metadata=block.metadata,
            )

        for block in parsed.footers:
            self._add_block(
                session,
                parsed.document_id,
                block_type="footer",
                text=block.text,
                metadata=block.metadata,
            )

        for comment in parsed.comments:
            self._add_block(
                session,
                parsed.document_id,
                block_type="comment",
                text=comment.text,
                metadata=comment.model_dump(exclude={"text"}),
            )

        for change in parsed.tracked_changes:
            self._add_block(
                session,
                parsed.document_id,
                block_type="tracked_change",
                text=change.text,
                metadata=change.model_dump(exclude={"text"}),
            )

        for hidden in parsed.hidden_text:
            self._add_block(
                session,
                parsed.document_id,
                block_type="hidden_text",
                text=hidden.text,
                metadata=hidden.metadata,
            )

        for table in parsed.tables:
            session.add(
                ParsedTable(
                    id=f"ptable_{uuid4().hex}",
                    document_id=parsed.document_id,
                    table_index=table.table_index,
                    caption=table.caption,
                    section_number=table.section_number,
                    section_title=table.section_title,
                    data_json=json.dumps(table.rows, ensure_ascii=False),
                    markdown=table.markdown,
                )
            )
            for row_index, row in enumerate(table.rows):
                for col_index, cell_text in enumerate(row):
                    self._add_block(
                        session,
                        parsed.document_id,
                        block_type="table_cell",
                        text=cell_text,
                        table_index=table.table_index,
                        row_index=row_index,
                        col_index=col_index,
                        section_number=table.section_number,
                        section_title=table.section_title,
                    )

        self._add_block(
            session,
            parsed.document_id,
            block_type="metadata",
            text=json.dumps(parsed.metadata, ensure_ascii=False),
            metadata=parsed.metadata,
        )

    def _add_block(
        self,
        session: Session,
        document_id: str,
        block_type: str,
        text: str,
        section_number: str | None = None,
        section_title: str | None = None,
        paragraph_index: int | None = None,
        table_index: int | None = None,
        row_index: int | None = None,
        col_index: int | None = None,
        style_name: str | None = None,
        metadata: dict | None = None,
    ) -> None:
        session.add(
            ParsedBlock(
                id=f"pblock_{uuid4().hex}",
                document_id=document_id,
                block_type=block_type,
                text=text,
                section_number=section_number,
                section_title=section_title,
                paragraph_index=paragraph_index,
                table_index=table_index,
                row_index=row_index,
                col_index=col_index,
                style_name=style_name,
                metadata_json=json.dumps(metadata or {}, ensure_ascii=False),
            )
        )

    def _metadata(self, doc: PythonDocxDocument) -> dict:
        props = doc.core_properties
        return {
            "title": props.title,
            "subject": props.subject,
            "author": props.author,
            "keywords": props.keywords,
            "comments": props.comments,
            "category": props.category,
            "created": props.created.isoformat() if props.created else None,
            "modified": props.modified.isoformat() if props.modified else None,
            "last_modified_by": props.last_modified_by,
            "revision": props.revision,
        }

    def _iter_body_blocks(self, doc: PythonDocxDocument):
        body = doc.element.body
        for child in body.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, doc)
            elif child.tag.endswith("}tbl"):
                yield DocxTable(child, doc)

    def _maybe_heading(self, text: str, style_name: str | None, paragraph_index: int) -> ParsedHeading | None:
        level = None
        if style_name and "heading" in style_name.lower():
            match = re.search(r"(\d+)$", style_name)
            level = int(match.group(1)) if match else None

        number = None
        title = text
        number_match = HEADING_NUMBER_RE.match(text)
        if number_match:
            number = number_match.group("number").rstrip(".")
            title = number_match.group("title").strip()
            level = level or number.count(".") + 1

        if level is None and number is None:
            return None

        return ParsedHeading(
            text=text,
            paragraph_index=paragraph_index,
            index=paragraph_index,
            style_name=style_name,
            section_number=number,
            section_title=title,
            level=level,
        )

    def _headers_and_footers(self, doc: PythonDocxDocument) -> tuple[list[ParsedTextBlock], list[ParsedTextBlock]]:
        headers: list[ParsedTextBlock] = []
        footers: list[ParsedTextBlock] = []
        for section_index, section in enumerate(doc.sections):
            for name in ("header", "first_page_header", "even_page_header"):
                text = self._paragraph_collection_text(getattr(section, name).paragraphs)
                if text:
                    headers.append(
                        ParsedTextBlock(
                            block_type="header",
                            text=text,
                            metadata={"section_index": section_index, "header_type": name},
                        )
                    )
            for name in ("footer", "first_page_footer", "even_page_footer"):
                text = self._paragraph_collection_text(getattr(section, name).paragraphs)
                if text:
                    footers.append(
                        ParsedTextBlock(
                            block_type="footer",
                            text=text,
                            metadata={"section_index": section_index, "footer_type": name},
                        )
                    )
        return headers, footers

    def _paragraph_collection_text(self, paragraphs: list[Paragraph]) -> str:
        return "\n".join(text for paragraph in paragraphs if (text := self._clean_text(paragraph.text)))

    def _comments(self, path: Path) -> list[ParsedComment]:
        xml = self._read_zip_xml(path, "word/comments.xml")
        if xml is None:
            return []
        comments = []
        for node in xml.xpath("//w:comment", namespaces=WORD_NS):
            text = self._node_text(node)
            if not text:
                continue
            comments.append(
                ParsedComment(
                    comment_id=node.get(f"{{{WORD_NS['w']}}}id"),
                    author=node.get(f"{{{WORD_NS['w']}}}author"),
                    date=node.get(f"{{{WORD_NS['w']}}}date"),
                    text=text,
                )
            )
        return comments

    def _tracked_changes(self, path: Path) -> list[TrackedChange]:
        xml = self._read_zip_xml(path, "word/document.xml")
        if xml is None:
            return []
        changes = []
        for tag in ("ins", "del", "moveFrom", "moveTo"):
            for node in xml.xpath(f"//w:{tag}", namespaces=WORD_NS):
                text = self._node_text(node)
                changes.append(
                    TrackedChange(
                        change_type=tag,
                        author=node.get(f"{{{WORD_NS['w']}}}author"),
                        date=node.get(f"{{{WORD_NS['w']}}}date"),
                        text=text,
                        xml_location_hint=f"word/document.xml//w:{tag}",
                    )
                )
        return changes

    def _hidden_text(self, path: Path) -> list[ParsedTextBlock]:
        xml = self._read_zip_xml(path, "word/document.xml")
        if xml is None:
            return []
        blocks = []
        for run in xml.xpath("//w:r[w:rPr/w:vanish]", namespaces=WORD_NS):
            text = self._node_text(run)
            if text:
                blocks.append(ParsedTextBlock(block_type="hidden_text", text=text, metadata={"source": "w:vanish"}))
        return blocks

    def _read_zip_xml(self, path: Path, member: str) -> etree._Element | None:
        try:
            with zipfile.ZipFile(path) as docx_zip:
                if member not in docx_zip.namelist():
                    return None
                return etree.fromstring(docx_zip.read(member))
        except (zipfile.BadZipFile, etree.XMLSyntaxError):
            return None

    def _node_text(self, node: etree._Element) -> str:
        return self._clean_text("".join(node.xpath(".//w:t/text() | .//w:delText/text()", namespaces=WORD_NS)))

    def _find_caption(self, candidates: list[str]) -> str | None:
        for text in reversed(candidates):
            if re.match(r"^\s*(Table|Figure)\s+\d+(?:-\d+)?[\.:]?", text, re.IGNORECASE) or re.match(
                r"^\s*表\s*\d+", text
            ):
                return text
        return None

    def _to_markdown(self, rows: list[list[str]]) -> str:
        if not rows:
            return ""
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        lines = ["| " + " | ".join(normalized[0]) + " |"]
        lines.append("| " + " | ".join(["---"] * width) + " |")
        for row in normalized[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    def _clean_text(self, text: str) -> str:
        return re.sub(r"[\r\n\t]+", " ", text.replace("\xa0", " ")).strip()
