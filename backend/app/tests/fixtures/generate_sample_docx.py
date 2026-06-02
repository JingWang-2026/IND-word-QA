from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

FIXTURE_DIR = Path(__file__).resolve().parent
BAD_DOCX = FIXTURE_DIR / "sample_bad_report.docx"
CLEAN_DOCX = FIXTURE_DIR / "sample_clean_report.docx"


def create_sample_bad_report(path: Path = BAD_DOCX) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Protocol Version 1.0"
    section.footer.paragraphs[0].text = "Protocol v1.1"

    doc.add_heading("3.1 Study Design", level=2)
    doc.add_paragraph("The the dose  escalation will follow the protocol.")
    doc.add_paragraph("3/20 subjects (20.0%) had treatment-emergent adverse events.")
    doc.add_paragraph("See Table 9 for additional disposition details.")
    doc.add_paragraph("AE was monitored during the study.")
    doc.add_paragraph("dose-limiting toxicity (DLT) was recorded.")
    doc.add_paragraph("dose level testing (DLT) was also recorded.")
    doc.add_paragraph("ABC-123 was supplied. ABC123 was tested.")
    doc.add_heading("3.3 Dose Escalation", level=2)

    doc.add_paragraph("Table 1. Subject disposition")
    table = doc.add_table(rows=5, cols=2)
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Count"
    table.cell(1, 0).text = "Male"
    table.cell(1, 1).text = "12"
    table.cell(2, 0).text = "Female"
    table.cell(2, 1).text = "9"
    table.cell(3, 0).text = "Unknown"
    table.cell(3, 1).text = ""
    table.cell(4, 0).text = "Total"
    table.cell(4, 1).text = "20"

    doc.add_paragraph("Table 2. Applicability")
    na_table = doc.add_table(rows=3, cols=1)
    na_table.cell(0, 0).text = "N/A"
    na_table.cell(1, 0).text = "NA"
    na_table.cell(2, 0).text = "Not Applicable"

    doc.save(path)
    _inject_review_xml(path)
    return path


def create_sample_clean_report(path: Path = CLEAN_DOCX) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Protocol Version 1.0"
    section.footer.paragraphs[0].text = "Protocol Version 1.0"

    doc.add_heading("1 Study Design", level=1)
    doc.add_paragraph("The dose escalation will follow the protocol.")
    doc.add_paragraph("3/20 subjects (15.0%) had treatment-emergent adverse events.")
    doc.add_paragraph("Table 1. Subject disposition")
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Count"
    table.cell(1, 0).text = "Male"
    table.cell(1, 1).text = "12"
    table.cell(2, 0).text = "Female"
    table.cell(2, 1).text = "8"
    table.cell(3, 0).text = "Total"
    table.cell(3, 1).text = "20"
    doc.add_paragraph("See Table 1 for details.")

    doc.save(path)
    return path


def _inject_review_xml(path: Path) -> None:
    with ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}

    members["word/comments.xml"] = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0" w:author="QA Reviewer" w:date="2026-06-01T10:00:00Z">
    <w:p><w:r><w:t>Remove this comment before finalization.</w:t></w:r></w:p>
  </w:comment>
</w:comments>"""

    document_xml = members["word/document.xml"].decode("utf-8")
    document_xml = document_xml.replace(
        "<w:body>",
        '<w:body><w:p><w:ins w:author="QA Reviewer" w:date="2026-06-01T10:00:00Z"><w:r><w:t>Inserted tracked text</w:t></w:r></w:ins></w:p>',
        1,
    )
    members["word/document.xml"] = document_xml.encode("utf-8")

    with ZipFile(path, "w", ZIP_DEFLATED) as target:
        for name, content in members.items():
            target.writestr(name, content)


if __name__ == "__main__":
    create_sample_bad_report()
    create_sample_clean_report()
    print(BAD_DOCX)
    print(CLEAN_DOCX)
