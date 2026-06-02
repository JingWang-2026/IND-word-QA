from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document as DocxDocument
from sqlmodel import Session, select

from app.core.database import engine, init_db
from app.models.parsed_content import ParsedBlock, ParsedTable
from app.services.docx_parser import DocxParserService


def test_docx_parser_extracts_core_content(tmp_path: Path) -> None:
    init_db()
    path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.core_properties.author = "QA Author"
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Protocol Version 1.0"
    section.footer.paragraphs[0].text = "Confidential"
    doc.add_heading("1 Study Design", level=1)
    doc.add_paragraph("The dose escalation will follow the protocol.")
    doc.add_paragraph("Table 1. Subject disposition")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Count"
    table.cell(1, 0).text = "Male"
    table.cell(1, 1).text = "12"
    table.cell(2, 0).text = "Total"
    table.cell(2, 1).text = "20"
    doc.save(path)

    parsed = DocxParserService().parse("doc_parser_test", str(path))
    assert len(parsed.paragraphs) >= 3
    assert parsed.headings[0].section_number == "1"
    assert parsed.tables[0].caption == "Table 1. Subject disposition"
    assert parsed.headers[0].text == "Protocol Version 1.0"
    assert parsed.footers[0].text == "Confidential"
    assert parsed.metadata["author"] == "QA Author"

    with Session(engine) as session:
        DocxParserService().save_to_database(session, parsed)
        session.commit()
        blocks = session.exec(select(ParsedBlock).where(ParsedBlock.document_id == "doc_parser_test")).all()
        tables = session.exec(select(ParsedTable).where(ParsedTable.document_id == "doc_parser_test")).all()

    assert any(block.block_type == "paragraph" for block in blocks)
    assert any(block.block_type == "header" for block in blocks)
    assert len(tables) == 1


def test_docx_parser_detects_comments_and_tracked_changes(tmp_path: Path) -> None:
    path = tmp_path / "raw_xml.docx"
    doc = DocxDocument()
    doc.add_paragraph("Original text.")
    doc.save(path)

    with ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}

    comments_xml = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0" w:author="Reviewer" w:date="2026-06-01T10:00:00Z">
    <w:p><w:r><w:t>Resolve this comment.</w:t></w:r></w:p>
  </w:comment>
</w:comments>"""
    document_xml = members["word/document.xml"].decode("utf-8")
    document_xml = document_xml.replace(
        "<w:body>",
        '<w:body><w:p><w:ins w:author="Reviewer" w:date="2026-06-01T10:00:00Z"><w:r><w:t>Inserted text</w:t></w:r></w:ins></w:p>',
        1,
    )
    members["word/comments.xml"] = comments_xml
    members["word/document.xml"] = document_xml.encode("utf-8")

    with ZipFile(path, "w", ZIP_DEFLATED) as target:
        for name, content in members.items():
            target.writestr(name, content)

    parsed = DocxParserService().parse("doc_raw_xml_test", str(path))
    assert parsed.comments[0].text == "Resolve this comment."
    assert parsed.tracked_changes[0].change_type == "ins"
    assert parsed.tracked_changes[0].text == "Inserted text"
